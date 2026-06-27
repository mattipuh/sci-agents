import io
import uuid
import logging
from dataclasses import dataclass

log = logging.getLogger(__name__)

_store: dict[str, "Document"] = {}


@dataclass
class Document:
    id: str
    filename: str
    text: str
    char_count: int


def extract_text(filename: str, content: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".txt") or lower.endswith(".md"):
        return content.decode("utf-8", errors="replace")
    if lower.endswith(".pdf"):
        return _extract_pdf(content)
    if lower.endswith(".docx"):
        return _extract_docx(content)
    raise ValueError(f"Unsupported file type: {filename}. Use .txt, .md, .pdf, or .docx")


def _extract_pdf(content: bytes) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def _extract_docx(content: bytes) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(content))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def store_document(filename: str, content: bytes) -> Document:
    text = extract_text(filename, content)
    doc_id = uuid.uuid4().hex[:12]
    doc = Document(id=doc_id, filename=filename, text=text, char_count=len(text))
    _store[doc_id] = doc
    log.info("Stored document %s (%s, %d chars)", doc_id, filename, doc.char_count)
    return doc


def get_document(doc_id: str) -> Document | None:
    return _store.get(doc_id)


def get_documents(doc_ids: list[str]) -> list[Document]:
    return [_store[did] for did in doc_ids if did in _store]


def clear_store():
    _store.clear()
